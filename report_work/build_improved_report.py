from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


OUT_DIR = Path(__file__).resolve().parent
OUT_PATH = OUT_DIR / "PFE_REPORT_improved.docx"
ASSET_DIR = OUT_DIR / "assets"
ASSET_DIR.mkdir(exist_ok=True)


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(33, 37, 41)
MUTED = RGBColor(86, 96, 106)
LIGHT_FILL = "F2F4F7"
ACCENT_FILL = "E8EEF5"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def paragraph_border_bottom(paragraph, color="BFBFBF") -> None:
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)


def set_run(run, *, size=None, bold=None, italic=None, color=None, font="Calibri") -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_paragraph(doc: Document, text: str = "", style: str | None = None, *, align=None):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        set_run(r, color=INK)
    if align is not None:
        p.alignment = align
    return p


def add_body(doc: Document, text: str) -> None:
    add_paragraph(doc, text, "Body Text")


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run(r, color=INK)


def add_numbers(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_run(r, color=INK)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], LIGHT_FILL)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hdr[i].paragraphs[0].add_run(h)
        set_run(run, bold=True, color=DARK_BLUE)
    for row_data in rows:
        row = table.add_row().cells
        for i, value in enumerate(row_data):
            p = row[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(value)
            set_run(run, size=10, color=INK)
    set_table_width(table, widths)
    doc.add_paragraph()


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run(r, size=9, italic=True, color=MUTED)


def add_toc(doc: Document) -> None:
    entries = [
        ("Dedication", 2),
        ("Acknowledgments", 3),
        ("Abstract", 4),
        ("Résumé", 5),
        ("List of Figures", 7),
        ("List of Tables", 8),
        ("Chapter 1 - Presentation of Oracle and the Project Environment", 9),
        ("Chapter 2 - Domain Background and Key Concepts: Remote Sensing", 11),
        ("Chapter 3 - General Project Context", 14),
        ("Chapter 4 - Conception and Architecture", 17),
        ("Chapter 5 - Implementation of the Proposed Solution", 19),
        ("Chapter 6 - Validation, Results and Discussion", 21),
        ("General Conclusion and Perspectives", 24),
        ("Webography", 24),
    ]
    for title, page in entries:
        p = doc.add_paragraph(style="Body Text")
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.25), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        r = p.add_run(f"{title}\t{page}")
        set_run(r, color=INK)


def create_pipeline_figure(path: Path) -> None:
    img = Image.new("RGB", (1400, 520), "white")
    draw = ImageDraw.Draw(img)
    try:
        font_title = ImageFont.truetype("Arial Bold.ttf", 34)
        font = ImageFont.truetype("Arial.ttf", 24)
        font_small = ImageFont.truetype("Arial.ttf", 19)
    except OSError:
        font_title = ImageFont.load_default()
        font = ImageFont.load_default()
        font_small = ImageFont.load_default()
    draw.text((55, 35), "NimbusChain end-to-end processing flow", fill=(31, 77, 120), font=font_title)
    stages = [
        ("AOI", "User area\nand filters"),
        ("Search", "Provider\npreview"),
        ("Download", "Copernicus\nor USGS"),
        ("Sen2Like", "Landsat to\nSentinel-like"),
        ("Zarr", "Scene-level\narrays"),
        ("Masks", "Cloud and\nwater layers"),
        ("Cube", "Optional time\ncube output"),
    ]
    x = 55
    y = 150
    w = 150
    h = 150
    gap = 38
    for idx, (title, subtitle) in enumerate(stages):
        fill = (232, 238, 245) if idx % 2 == 0 else (242, 244, 247)
        draw.rounded_rectangle((x, y, x + w, y + h), radius=18, fill=fill, outline=(46, 116, 181), width=3)
        draw.text((x + 18, y + 25), title, fill=(31, 77, 120), font=font)
        draw.multiline_text((x + 18, y + 72), subtitle, fill=(33, 37, 41), font=font_small, spacing=6)
        if idx < len(stages) - 1:
            draw.line((x + w + 6, y + h // 2, x + w + gap - 8, y + h // 2), fill=(46, 116, 181), width=4)
            draw.polygon(
                [
                    (x + w + gap - 8, y + h // 2),
                    (x + w + gap - 26, y + h // 2 - 10),
                    (x + w + gap - 26, y + h // 2 + 10),
                ],
                fill=(46, 116, 181),
            )
        x += w + gap
    draw.text(
        (55, 375),
        "The platform keeps one job identifier across the complete lifecycle: discovery, acquisition, conversion, masking and result exposure.",
        fill=(86, 96, 106),
        font=font_small,
    )
    img.save(path)


def create_architecture_figure(path: Path) -> None:
    img = Image.new("RGB", (1400, 840), "white")
    draw = ImageDraw.Draw(img)
    try:
        title = ImageFont.truetype("Arial Bold.ttf", 34)
        font = ImageFont.truetype("Arial.ttf", 24)
        bold = ImageFont.truetype("Arial Bold.ttf", 24)
        small = ImageFont.truetype("Arial.ttf", 19)
    except OSError:
        title = font = bold = small = ImageFont.load_default()
    draw.text((55, 35), "Service-oriented architecture", fill=(31, 77, 120), font=title)
    boxes = {
        "UI": (85, 150, 365, 280, "Streamlit UI", "AOI selection, preview,\njob monitoring"),
        "API": (560, 150, 840, 280, "FastAPI API", "Validation, jobs,\nartifacts, health"),
        "WORKER": (1035, 150, 1315, 280, "Worker", "Downloads and pipeline\norchestration"),
        "ZARR": (330, 450, 610, 580, "Zarr service", "Normalization,\nbands, cubes"),
        "MASK": (790, 450, 1070, 580, "Mask service", "Cloud probability,\nwater inference"),
        "STORE": (535, 620, 875, 735, "Shared state/storage", "SQLite/MongoDB, downloads,\nZarr artifacts"),
    }
    for key, (x1, y1, x2, y2, name, sub) in boxes.items():
        fill = (232, 238, 245) if key in {"API", "STORE"} else (242, 244, 247)
        draw.rounded_rectangle((x1, y1, x2, y2), radius=16, fill=fill, outline=(46, 116, 181), width=3)
        draw.text((x1 + 20, y1 + 22), name, fill=(31, 77, 120), font=bold)
        draw.multiline_text((x1 + 20, y1 + 62), sub, fill=(33, 37, 41), font=small, spacing=6)

    arrows = [
        ((365, 215), (560, 215)),
        ((840, 215), (1035, 215)),
        ((1135, 280), (930, 450)),
        ((1135, 280), (470, 450)),
        ((610, 515), (790, 515)),
        ((700, 580), (700, 620)),
    ]
    for start, end in arrows:
        draw.line((start[0], start[1], end[0], end[1]), fill=(46, 116, 181), width=4)
        ex, ey = end
        sx, sy = start
        if abs(ex - sx) > abs(ey - sy):
            direction = 1 if ex > sx else -1
            draw.polygon([(ex, ey), (ex - 18 * direction, ey - 9), (ex - 18 * direction, ey + 9)], fill=(46, 116, 181))
        else:
            direction = 1 if ey > sy else -1
            draw.polygon([(ex, ey), (ex - 9, ey - 18 * direction), (ex + 9, ey - 18 * direction)], fill=(46, 116, 181))
    draw.text((55, 785), "Heavy operations are isolated from the API so long-running jobs do not block the user interface.", fill=(86, 96, 106), font=small)
    img.save(path)


def create_water_solution_figure(path: Path) -> None:
    img = Image.new("RGB", (1400, 520), "white")
    draw = ImageDraw.Draw(img)
    try:
        title = ImageFont.truetype("Arial Bold.ttf", 34)
        font = ImageFont.truetype("Arial.ttf", 23)
        bold = ImageFont.truetype("Arial Bold.ttf", 24)
        small = ImageFont.truetype("Arial.ttf", 19)
    except OSError:
        title = font = bold = small = ImageFont.load_default()
    draw.text((55, 35), "Recommended water-mask runtime placement", fill=(31, 77, 120), font=title)
    cols = [
        (70, 150, 400, 365, "Current local Podman", "OmniWater runs on CPU only.\nPrecise, but too slow on full scenes."),
        (535, 150, 865, 365, "Preferred accelerator path", "Remote OCI GPU service or\nhost-native Apple MPS runner."),
        (1000, 150, 1330, 365, "CPU fallback", "Small scenes only with 4096 px\nmodel exports and 512 px outputs."),
    ]
    for x1, y1, x2, y2, name, sub in cols:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=18, fill=(242, 244, 247), outline=(46, 116, 181), width=3)
        draw.text((x1 + 22, y1 + 28), name, fill=(31, 77, 120), font=bold)
        draw.multiline_text((x1 + 22, y1 + 82), sub, fill=(33, 37, 41), font=small, spacing=8)
    draw.line((400, 257, 535, 257), fill=(46, 116, 181), width=5)
    draw.polygon([(535, 257), (512, 246), (512, 268)], fill=(46, 116, 181))
    draw.line((865, 257, 1000, 257), fill=(46, 116, 181), width=5)
    draw.polygon([(1000, 257), (977, 246), (977, 268)], fill=(46, 116, 181))
    draw.text(
        (70, 425),
        "The model logic remains unchanged; the fix is to move inference to a runtime that can actually expose GPU/MPS acceleration.",
        fill=(86, 96, 106),
        font=font,
    )
    img.save(path)


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    body = styles["Body Text"]
    body.font.name = "Calibri"
    body._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    body.font.size = Pt(11)
    body.font.color.rgb = INK
    body.paragraph_format.space_after = Pt(6)
    body.paragraph_format.line_spacing = 1.10
    body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = MUTED
    caption.paragraph_format.space_after = Pt(8)

    for list_name in ("List Bullet", "List Number"):
        st = styles[list_name]
        st.font.name = "Calibri"
        st.font.size = Pt(11)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.167

    if "Callout" not in styles:
        callout = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        callout = styles["Callout"]
    callout.font.name = "Calibri"
    callout.font.size = Pt(10.5)
    callout.font.color.rgb = INK
    callout.paragraph_format.left_indent = Inches(0.18)
    callout.paragraph_format.right_indent = Inches(0.18)
    callout.paragraph_format.space_before = Pt(6)
    callout.paragraph_format.space_after = Pt(8)
    callout.paragraph_format.line_spacing = 1.10

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("NimbusChain Pipeline - End-of-Studies Project Report")
    set_run(run, size=9, color=MUTED)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("End-of-Studies Project Report")
    set_run(r, size=22, bold=True, color=DARK_BLUE)
    paragraph_border_bottom(p)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Design and Implementation of an End-to-End Satellite Imagery Acquisition and Processing Pipeline")
    set_run(r, size=18, bold=True, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("NimbusChain Pipeline")
    set_run(r, size=16, bold=True, color=BLUE)

    doc.add_paragraph()
    for line in [
        "Submitted in partial fulfillment of the requirements for the Software Engineering Degree",
        "Major: Computer Engineering",
        "Promotion 2025 / 2026",
        "Internship carried out within Oracle Morocco R&D Center - Digital Government",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        set_run(r, size=12, color=INK)

    doc.add_paragraph()
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Written by", "Mr. Mehdi Dinari"),
        ("Company mentor", "Mr. Hamza El Makrini"),
        ("School tutor", "Mr. Abdelhak Boulaalam"),
        ("Defense date", "................................"),
    ]
    for row, (left, right) in zip(table.rows, rows):
        row.cells[0].text = left
        row.cells[1].text = right
        set_cell_shading(row.cells[0], LIGHT_FILL)
        for cell in row.cells:
            set_cell_margins(cell)
            for para in cell.paragraphs:
                for run in para.runs:
                    set_run(run, color=INK)
    set_table_width(table, [2900, 6460])
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Jury Members: Mr. Hamza El Makrini, Mr. Abdelhak Boulaalam, and invited ENSAF professors")
    set_run(r, size=11, color=MUTED)
    doc.add_page_break()


def add_front_matter(doc: Document) -> None:
    doc.add_heading("Dedication", level=1)
    add_body(
        doc,
        "This work is dedicated to my family, my supervisors and my friends, whose support, patience and encouragement accompanied me throughout this internship and engineering journey.",
    )
    add_body(
        doc,
        "Their trust gave me the motivation to overcome technical difficulties, learn from each iteration and complete this project with discipline and curiosity.",
    )
    doc.add_page_break()

    doc.add_heading("Acknowledgments", level=1)
    add_body(
        doc,
        "I would like to express my sincere gratitude to my company mentor, Mr. Hamza El Makrini, for his guidance, availability and technical feedback throughout this internship. His advice helped me understand the project environment, improve the quality of my work and approach engineering decisions with more rigor.",
    )
    add_body(
        doc,
        "I also thank my school tutor, Mr. Abdelhak Boulaalam, for his academic guidance and constructive feedback. His support contributed to the organization of this report and to the alignment of the project with the expectations of an end-of-studies engineering work.",
    )
    add_body(
        doc,
        "My sincere thanks go to my managers, Mr. Faical Tabouti and Mr. Sam Batterman, for the supportive environment they provided and for their availability whenever I needed assistance or clarification.",
    )
    add_body(
        doc,
        "I extend my appreciation to the Oracle Digital Government team and to the administration, professors and staff of ENSA Fès for the knowledge, professional values and learning conditions they provided during my training.",
    )
    add_body(
        doc,
        "Finally, I thank the members of the jury for accepting to evaluate this work, as well as my family and friends for their constant encouragement.",
    )
    doc.add_page_break()

    doc.add_heading("Abstract", level=1)
    add_body(
        doc,
        "This report presents the design and implementation of an end-to-end satellite imagery acquisition and processing pipeline developed during an internship within Oracle's Digital Government context. The project addresses a practical operational need: analysts need to define an area of interest, preview available scenes, launch provider downloads, normalize the outputs into analysis-ready structures and monitor the full execution lifecycle from a single interface.",
    )
    add_body(
        doc,
        "The implemented platform is service-oriented. It combines a FastAPI backend, background workers, a Streamlit user interface, a Zarr conversion service and a mask service for cloud and water processing. The pipeline supports provider preview, asynchronous job orchestration, provider-aware download coordination, conversion of raw scenes to normalized scene-level Zarr stores, optional cube generation and execution in local or cloud environments.",
    )
    add_body(
        doc,
        "The resulting system improves traceability, reduces manual intervention and provides a reusable basis for operational Earth-observation workflows. The validation phase also highlights a key engineering constraint: precise water masking with OmniWater requires an accelerator-backed runtime for full scenes, while the CPU fallback remains useful only for smaller workloads.",
    )
    add_body(
        doc,
        "Keywords: Remote Sensing, Satellite Imagery, Zarr, Microservices, Cloud Mask, Water Mask, Oracle Cloud Infrastructure.",
    )
    doc.add_page_break()

    doc.add_heading("Résumé", level=1)
    add_body(
        doc,
        "Ce rapport présente la conception et la mise en oeuvre d'une plateforme de téléchargement et de traitement d'images satellitaires développée dans le cadre d'un stage au sein du contexte Digital Government d'Oracle. Le projet répond à un besoin opérationnel concret: permettre à un utilisateur de définir une zone d'intérêt, de prévisualiser les scènes disponibles, de lancer les téléchargements depuis plusieurs fournisseurs, de normaliser les sorties dans un format exploitable et de suivre tout le cycle d'exécution depuis une interface unique.",
    )
    add_body(
        doc,
        "La solution proposée repose sur une architecture orientée services composée d'un backend FastAPI, de workers d'arrière-plan, d'une interface Streamlit, d'un service de conversion vers Zarr et d'un service de masquage nuage/eau. La plateforme prend en charge la prévisualisation des produits, l'orchestration asynchrone des jobs, la coordination des téléchargements, la conversion des scènes brutes vers des Zarr normalisés, la construction optionnelle de cubes temporels et l'exécution en local ou sur Oracle Cloud Infrastructure.",
    )
    add_body(
        doc,
        "Les résultats montrent que l'architecture améliore la traçabilité, limite les manipulations manuelles et facilite la réutilisation des sorties. Le rapport met également en évidence une limite importante: pour obtenir un masque d'eau précis et rapide sur des scènes complètes, le modèle OmniWater doit être exécuté sur un runtime accéléré, tandis que le mode CPU doit rester un mécanisme de secours.",
    )
    add_body(
        doc,
        "Mots-clés: télédétection, imagerie satellitaire, Zarr, microservices, masque nuage, masque eau, Oracle Cloud Infrastructure.",
    )
    doc.add_page_break()

    doc.add_heading("Table of Contents", level=1)
    add_toc(doc)
    doc.add_page_break()

    doc.add_heading("List of Figures", level=1)
    add_body(doc, "Figure 2.1: NimbusChain end-to-end processing flow.")
    add_body(doc, "Figure 3.1: Service-oriented architecture.")
    add_body(doc, "Figure 6.1: Recommended water-mask runtime placement.")
    doc.add_page_break()

    doc.add_heading("List of Tables", level=1)
    add_body(doc, "Table 2.1: Main satellite product families handled by the pipeline.")
    add_body(doc, "Table 2.2: Spectral band categories used in remote sensing interpretation.")
    add_body(doc, "Table 3.1: High-level architecture of the platform.")
    add_body(doc, "Table 5.1: Main technology stack used by the project.")
    add_body(doc, "Table 6.1: Local validation observations.")
    doc.add_page_break()


def add_chapter_1(doc: Document) -> None:
    doc.add_heading("Chapter 1 - Presentation of Oracle and the Project Environment", level=1)
    doc.add_heading("Introduction", level=2)
    add_body(
        doc,
        "This chapter presents the professional environment in which the internship was carried out. It introduces Oracle Corporation, the Oracle Morocco Research and Development Center, the Digital Government group and the organizational context that framed the project.",
    )
    doc.add_heading("Oracle Corporation", level=2)
    add_body(
        doc,
        "Oracle Corporation is a multinational technology company founded in 1977 by Larry Ellison, Bob Miner and Ed Oates. The company first became known through Oracle Database, a relational database management system that became one of the foundations of modern enterprise information systems.",
    )
    add_body(
        doc,
        "Over time, Oracle expanded its portfolio to include enterprise applications, cloud infrastructure, middleware, hardware systems, developer tools, consulting services and advanced data platforms. The acquisition of Sun Microsystems in 2010 also strengthened Oracle's position in major technologies such as Java and Solaris.",
    )
    doc.add_heading("Oracle Morocco R&D Center", level=2)
    add_body(
        doc,
        "The Oracle Morocco Research and Development Center is one of Oracle's strategic engineering centers in the EMEA region. Located in Casablanca, it brings together software engineers, data specialists, project managers and research-oriented teams working on a broad range of Oracle products and initiatives.",
    )
    add_body(
        doc,
        "The center contributes to domains such as cloud computing, data platforms, artificial intelligence, cybersecurity, analytics and digital public services. It also supports internships, graduate recruitment initiatives and collaboration with Moroccan engineering schools.",
    )
    doc.add_heading("Digital Government Context", level=2)
    add_body(
        doc,
        "The project was carried out within the Digital Government context. This environment focuses on digital platforms that help public-sector actors use data more effectively, improve service delivery and support operational decision-making.",
    )
    add_body(
        doc,
        "Geospatial information is particularly relevant in this context because it can support agriculture, land-use observation, water monitoring, infrastructure planning and environmental reporting. However, transforming satellite products into reliable digital services requires more than downloading images. It requires reproducible pipelines, runtime visibility, normalized outputs and deployment strategies adapted to heavy geospatial workloads.",
    )
    doc.add_heading("Project Motivation", level=2)
    add_body(
        doc,
        "The internship project addresses the need for a reusable satellite-processing platform. Existing workflows are often fragmented across provider portals, local scripts and manual processing steps. This fragmentation makes it difficult to follow progress, recover from failures, preserve output lineage and expose the workflow to non-expert users.",
    )
    add_body(
        doc,
        "The goal of the platform is therefore to transform a scattered workflow into an operational chain that supports product discovery, asynchronous execution, normalized Zarr outputs, optional cube building and cloud or water masking.",
    )
    doc.add_heading("Conclusion", level=2)
    add_body(
        doc,
        "This chapter introduced Oracle and the professional context of the internship. It also clarified why satellite imagery processing is relevant for Digital Government use cases and why an operational pipeline is needed before downstream geospatial analysis can be performed reliably.",
    )


def add_chapter_2(doc: Document) -> None:
    doc.add_heading("Chapter 2 - Domain Background and Key Concepts: Remote Sensing", level=1)
    doc.add_heading("Introduction", level=2)
    add_body(
        doc,
        "This chapter introduces the remote sensing concepts required to understand the NimbusChain pipeline. The objective is not to provide an exhaustive scientific presentation of Earth observation, but to explain the data structures, satellite missions and preprocessing needs that directly influenced the engineering design.",
    )
    doc.add_heading("Earth Observation Context", level=2)
    add_body(
        doc,
        "Earth observation consists of acquiring information about the Earth's surface using remote sensing technologies, especially satellites. For agricultural and environmental applications, satellite imagery makes it possible to observe large areas repeatedly and objectively.",
    )
    add_body(
        doc,
        "The platform focuses on workflows where users select an area of interest, choose dates and providers, inspect candidate products and transform raw scenes into outputs that are easier to analyze. This requires both geospatial understanding and robust software orchestration.",
    )
    doc.add_heading("Satellite Product Families", level=2)
    add_body(
        doc,
        "The project mainly considers Sentinel-2 and Landsat 8/9 products. These missions are widely used in optical Earth observation and are relevant for vegetation monitoring, land-cover analysis, water observation and environmental reporting.",
    )
    add_table(
        doc,
        ["Mission", "Main characteristics", "Role in the project"],
        [
            ["Sentinel-2", "Multispectral optical mission with bands at 10 m, 20 m and 60 m resolution.", "Used as a reference target structure for high-resolution optical analysis and Zarr normalization."],
            ["Landsat 8/9", "Long-running Earth observation program with optical, panchromatic and thermal information.", "Used as an institutional provider source, especially through USGS access and Sen2Like harmonization."],
            ["Sen2Like output", "Processing output that makes Landsat products closer to Sentinel-like structures.", "Allows downstream stages to rely on a more homogeneous representation across providers."],
        ],
        [1800, 3900, 3660],
    )
    add_caption(doc, "Table 2.1: Main satellite product families handled by the pipeline.")
    doc.add_heading("Spectral Bands and Interpretation", level=2)
    add_body(
        doc,
        "Satellite sensors capture information in several spectral bands. Unlike ordinary RGB images, multispectral products include visible, near-infrared, shortwave infrared and sometimes thermal bands. These bands carry complementary information about vegetation, water, soil, clouds and surface conditions.",
    )
    add_table(
        doc,
        ["Band category", "Interpretation", "Typical downstream use"],
        [
            ["Visible bands", "Blue, green and red bands approximate human visual perception.", "RGB preview, visual inspection, land-cover distinction."],
            ["Near-infrared", "Healthy vegetation strongly reflects NIR energy.", "Vegetation condition, biomass and crop vigor analysis."],
            ["Red-edge", "Sensitive to subtle vegetation and chlorophyll variations.", "Agriculture-oriented monitoring and stress detection."],
            ["SWIR", "Useful for moisture, soil and vegetation condition analysis.", "Water-related indices, surface moisture and stress analysis."],
            ["Quality layers", "Provider-generated or model-generated scene quality information.", "Cloud, shadow, snow, water or invalid-pixel filtering."],
        ],
        [1900, 4000, 3460],
    )
    add_caption(doc, "Table 2.2: Spectral band categories used in remote sensing interpretation.")
    doc.add_heading("Tile Systems and Spatial Organization", level=2)
    add_body(
        doc,
        "Satellite products are not distributed as one continuous global image. Sentinel-2 uses the MGRS tiling system, while Landsat products are organized through the WRS-2 path/row logic. A single user area may therefore intersect one or several tiles, which affects search, download, conversion and cube generation.",
    )
    doc.add_heading("Why Preprocessing Is Needed", level=2)
    add_body(
        doc,
        "Raw products are heterogeneous. They differ in file naming, metadata, spatial resolution, band organization and provider-specific conventions. They may also include clouds, shadows or other artifacts that reduce direct analytical value.",
    )
    add_body(
        doc,
        "Preprocessing in NimbusChain therefore includes provider-aware download coordination, Sen2Like normalization when needed, conversion to Zarr, optional cloud and water masking and optional temporal cube building. These stages transform raw provider outputs into structured artifacts that can be reused by later analytical services.",
    )
    create_pipeline_figure(ASSET_DIR / "pipeline_flow.png")
    doc.add_picture(str(ASSET_DIR / "pipeline_flow.png"), width=Inches(6.5))
    add_caption(doc, "Figure 2.1: NimbusChain end-to-end processing flow.")
    doc.add_heading("Conclusion", level=2)
    add_body(
        doc,
        "Remote sensing data is rich, but operationally complex. Understanding missions, bands, tile systems and preprocessing requirements explains why the project needs a modular pipeline rather than a simple downloader.",
    )


def add_chapter_3(doc: Document) -> None:
    doc.add_heading("Chapter 3 - General Project Context", level=1)
    doc.add_heading("Project Description", level=2)
    add_body(
        doc,
        "NimbusChain is an end-to-end satellite imagery acquisition and preparation platform. Its role is to move from user intent to analysis-ready artifacts through a traceable workflow: area selection, provider preview, job submission, download, normalization, masking and result exposure.",
    )
    add_body(
        doc,
        "The project is not limited to downloading data. Each stage produces structured outputs that can be inspected, persisted and reused. This is important because operational Earth-observation workflows often require several scenes, several providers and several downstream processing steps.",
    )
    doc.add_heading("Problem Statement", level=2)
    add_body(
        doc,
        "The initial problem was the absence of a unified operational pipeline able to manage satellite imagery from preview to analysis-ready outputs. Existing approaches depended on fragmented tools, manual provider interactions and ad hoc scripts. This made it difficult to track progress, recover from interruptions, control throughput and preserve lineage.",
    )
    doc.add_heading("Project Objectives", level=2)
    add_bullets(
        doc,
        [
            "Provide a user interface for AOI definition, product preview and job monitoring.",
            "Automate the transition from product selection to processed outputs.",
            "Execute long-running operations asynchronously through background workers.",
            "Normalize downloaded products into structured Zarr stores.",
            "Preserve satellite bands, metadata and auxiliary layers required by masks and analysis.",
            "Support cloud and water mask generation as explicit pipeline stages.",
            "Prepare outputs for optional temporal cube construction.",
            "Keep the codebase modular and extensible for future providers and services.",
        ],
    )
    doc.add_heading("Architecture Overview", level=2)
    add_body(
        doc,
        "The architecture follows a service-oriented design. The user interface handles interaction; the API handles validation, state and public endpoints; the worker executes long-running tasks; the Zarr service performs conversion and cube operations; and the mask service enriches generated Zarr stores with quality layers.",
    )
    add_table(
        doc,
        ["Component", "Technology", "Main responsibility"],
        [
            ["User interface", "Streamlit", "Collect inputs, preview products, submit jobs and display execution status."],
            ["Public API", "FastAPI", "Expose job, artifact, health, provider, conversion and masking endpoints."],
            ["Worker runtime", "Python asynchronous worker", "Claim jobs, coordinate downloads and drive the pipeline lifecycle."],
            ["Zarr service", "FastAPI service", "Convert raw scenes to Zarr and build optional cubes."],
            ["Mask service", "FastAPI service", "Apply cloud and water masks on existing Zarr outputs."],
            ["Persistence", "SQLite or MongoDB", "Store jobs, events, artifacts, state transitions and worker heartbeats."],
        ],
        [1900, 2200, 5260],
    )
    add_caption(doc, "Table 3.1: High-level architecture of the platform.")
    create_architecture_figure(ASSET_DIR / "service_architecture.png")
    doc.add_picture(str(ASSET_DIR / "service_architecture.png"), width=Inches(6.5))
    add_caption(doc, "Figure 3.1: Service-oriented architecture.")
    doc.add_heading("Why Modularity Is Essential", level=2)
    add_body(
        doc,
        "Modularity is a practical requirement for this project. The platform handles user interaction, provider communication, job orchestration, geospatial conversion, inference-based masking and artifact management. Mixing these responsibilities into one monolithic script would make the system fragile and difficult to evolve.",
    )
    add_body(
        doc,
        "By separating services, the platform improves maintainability, fault isolation, testability and deployment flexibility. It also makes it possible to move selected heavy stages, such as water masking, to a specialized runtime without redesigning the entire system.",
    )
    doc.add_heading("Conclusion", level=2)
    add_body(
        doc,
        "This chapter described the project as a modular satellite data pipeline platform. The next chapter focuses on the conception choices that organize these requirements into a coherent technical design.",
    )


def add_chapter_4(doc: Document) -> None:
    doc.add_heading("Chapter 4 - Conception and Architecture", level=1)
    doc.add_heading("Needs Analysis", level=2)
    add_body(
        doc,
        "The main user need is to launch and monitor a complete satellite-processing workflow without manually coordinating provider portals, local scripts and post-processing tools. The system must expose a clear operational surface while preserving enough technical detail for debugging and validation.",
    )
    doc.add_heading("Functional Requirements", level=2)
    add_bullets(
        doc,
        [
            "Search and preview satellite products from supported providers.",
            "Accept an area of interest, date range, provider and product configuration.",
            "Submit asynchronous jobs and expose their state, events and artifacts.",
            "Download selected products while respecting provider constraints.",
            "Run Sen2Like when Landsat harmonization is required.",
            "Convert raw or normalized scenes into Zarr stores with preserved bands and metadata.",
            "Apply cloud and water masks when requested.",
            "Build scene-level or grouped time cubes when the selected workflow requires it.",
        ],
    )
    doc.add_heading("Non-Functional Requirements", level=2)
    add_bullets(
        doc,
        [
            "Modularity: each runtime component must have a clear responsibility.",
            "Scalability: heavy execution should be isolated from the UI and API.",
            "Observability: jobs must expose progress, stage timings, errors and artifacts.",
            "Reliability: partial failures must be visible and recoverable.",
            "Portability: the stack must support local development and cloud deployment.",
            "Extensibility: new providers, masks and output formats should be addable without rewriting the core.",
        ],
    )
    doc.add_heading("Control Plane and Execution Plane", level=2)
    add_body(
        doc,
        "The control plane is represented by the FastAPI service. It receives user intent, validates requests, creates jobs, exposes state and keeps the UI responsive. The execution plane is represented by workers and processing services that perform downloads, conversions and inference-heavy tasks.",
    )
    add_body(
        doc,
        "This separation is important because satellite workflows are long-running by nature. A download or mask generation stage can take several minutes, and the API must remain available during that time.",
    )
    doc.add_heading("Data Model and Artifact Lineage", level=2)
    add_body(
        doc,
        "The platform relies on persistent job records and artifacts. Instead of returning large raster outputs directly through HTTP payloads, services exchange references such as job identifiers, source URIs, output Zarr URIs, scene identifiers and acquisition timestamps.",
    )
    add_body(
        doc,
        "This design reduces coupling between services and allows the UI to display a coherent timeline of the pipeline, including generated outputs and stage-level errors.",
    )
    doc.add_heading("Placement of Masking in the Workflow", level=2)
    add_body(
        doc,
        "Masking is designed as a post-conversion stage that enriches an existing Zarr store. This choice keeps raw acquisition, data normalization and mask inference separated. It also allows cloud and water masks to be recomputed or skipped without repeating the full download stage.",
    )
    add_body(
        doc,
        "The water mask stage has a specific runtime constraint: the accurate OmniWater model is significantly heavier than simple threshold-based approaches. For full-scene execution, the design must therefore support an accelerator-backed placement while keeping a CPU fallback for small scenes and tests.",
    )
    doc.add_heading("Conclusion", level=2)
    add_body(
        doc,
        "The conception phase defines the project as a traceable multi-service workflow. This design makes the pipeline easier to operate and prepares it for later optimization, especially for compute-heavy stages such as water masking.",
    )


def add_chapter_5(doc: Document) -> None:
    doc.add_heading("Chapter 5 - Implementation of the Proposed Solution", level=1)
    doc.add_heading("Implementation of the User Interface", level=2)
    add_body(
        doc,
        "The user interface is implemented with Streamlit. It provides the operational entry point for preparing requests, selecting provider parameters, previewing candidate products, submitting jobs and inspecting results. The UI acts as a client of the backend API and does not perform heavy processing directly.",
    )
    doc.add_heading("Implementation of the API Layer", level=2)
    add_body(
        doc,
        "The API layer is implemented with FastAPI. It exposes endpoints for job creation, job inspection, provider health, worker status, artifact discovery, conversion requests and mask operations. It centralizes validation and keeps the contract between UI and execution services stable.",
    )
    doc.add_heading("Worker and Job Execution Logic", level=2)
    add_body(
        doc,
        "The worker is responsible for claiming queued jobs and executing the pipeline. It records events, updates state transitions, coordinates downloads and calls downstream services. This makes the processing lifecycle persistent and observable even when individual stages are long-running.",
    )
    doc.add_heading("Zarr Conversion Workflow", level=2)
    add_body(
        doc,
        "The Zarr service converts raw or normalized products into structured arrays. The conversion preserves imagery bands, auxiliary layers, spatial context and provenance metadata so that downstream services can access the data consistently. Zarr was selected because its chunked multidimensional structure is suitable for large raster data and cloud-native access patterns.",
    )
    doc.add_heading("Masking Workflow", level=2)
    add_body(
        doc,
        "The mask service applies cloud and water masks on existing Zarr stores. Cloud masking produces both binary cloud layers and probability information when the selected backend supports it. Water masking uses OmniWater for model-based inference when the dependency and runtime are available.",
    )
    add_body(
        doc,
        "A key implementation detail is the separation between output resolution and model inference export size. The platform can keep 512 x 512 output chunks while exporting larger model tiles for OmniWater. This reduces excessive per-tile overhead without changing the final Zarr chunking strategy.",
    )
    doc.add_heading("Technology Stack", level=2)
    add_table(
        doc,
        ["Category", "Technology", "Role"],
        [
            ["Backend", "FastAPI, Uvicorn", "API endpoints, service contracts and health checks."],
            ["Frontend", "Streamlit", "Interactive workflow preparation and monitoring."],
            ["Execution", "Python async workers", "Background job processing and orchestration."],
            ["Geospatial", "Rasterio, Xarray, Zarr, PyProj", "Raster reading, array representation and spatial metadata handling."],
            ["Data validation", "Pydantic", "Typed request, response and configuration models."],
            ["Persistence", "SQLite or MongoDB", "Job records, events, artifacts and worker state."],
            ["Masking", "OmniCloudMask, OmniWaterMask", "Model-based cloud and water mask generation."],
            ["Deployment", "Docker, Podman, OCI", "Local container stack and cloud-oriented execution."],
        ],
        [1800, 2800, 4760],
    )
    add_caption(doc, "Table 5.1: Main technology stack used by the project.")
    doc.add_heading("Artifact Management", level=2)
    add_body(
        doc,
        "Each major stage registers artifacts so that results can be inspected after execution. Raw downloads, Sen2Like outputs, scene-level Zarr stores, mask layers and cubes are linked to the same job lifecycle. This improves traceability and makes debugging more efficient.",
    )
    doc.add_heading("Conclusion", level=2)
    add_body(
        doc,
        "The implementation translates the conceptual architecture into concrete services and workflows. The result is a modular platform where acquisition, conversion, masking and monitoring can evolve independently while remaining connected by shared contracts.",
    )


def add_chapter_6(doc: Document) -> None:
    doc.add_heading("Chapter 6 - Validation, Results and Discussion", level=1)
    doc.add_heading("Validation Strategy", level=2)
    add_body(
        doc,
        "The platform was validated through iterative end-to-end tests. Each test follows the same operational chain: provider search, download, Sen2Like normalization when required, Zarr conversion, optional cloud mask, optional water mask and final artifact inspection.",
    )
    doc.add_heading("Observed Local Results", level=2)
    add_body(
        doc,
        "Local validation showed that search, download, Sen2Like, Zarr conversion and cloud masking can execute through the stack. The water mask stage exposed the most important runtime limitation: in the current local Podman environment, the OmniWater model runs on CPU only, which makes full-scene inference too slow for comfortable interactive testing.",
    )
    add_table(
        doc,
        ["Stage", "Observed behavior", "Interpretation"],
        [
            ["Search", "Completed quickly during local tests.", "Provider metadata discovery is not the bottleneck."],
            ["Download", "Depends mainly on provider speed and scene size.", "Large Landsat products naturally dominate network and disk time."],
            ["Sen2Like", "Works but remains significant on full scenes.", "The stage should not be duplicated and should feed downstream Zarr directly."],
            ["Zarr conversion", "Completed in a reasonable time after normalization.", "Preserving bands and metadata is more important than aggressive simplification."],
            ["Cloud mask", "Acceptable runtime and useful output layers.", "The cloud path can remain local for many test scenarios."],
            ["Water mask", "Accurate model is slow on CPU-only Podman.", "Needs accelerator placement for precise and fast full-scene execution."],
        ],
        [1600, 3500, 4260],
    )
    add_caption(doc, "Table 6.1: Local validation observations.")
    doc.add_heading("Water Mask Issue and Proposed Solution", level=2)
    add_body(
        doc,
        "The water mask problem is not caused by Sen2Like and should not be solved by replacing the model with a heuristic. The accurate OmniWater path is model-based, but the current local container runtime does not expose a usable accelerator. As a result, the model falls back to CPU execution, which is precise but too slow on complete scenes.",
    )
    add_body(
        doc,
        "The recommended solution is to keep the OmniWater model logic unchanged and move its execution to an accelerator-backed runtime. Two practical placements are possible: a remote OCI GPU mask service connected to the same pipeline contracts, or a host-native macOS Apple MPS runner if the dependencies are installed outside Podman and shared Zarr paths are available.",
    )
    add_body(
        doc,
        "The CPU mode should remain a fallback. It can use large model export tiles, for example 4096 pixels, while keeping final 512 x 512 output chunks. This avoids generating hundreds of tiny model inputs, but it cannot replace GPU or MPS acceleration for full operational scenes.",
    )
    create_water_solution_figure(ASSET_DIR / "water_solution.png")
    doc.add_picture(str(ASSET_DIR / "water_solution.png"), width=Inches(6.5))
    add_caption(doc, "Figure 6.1: Recommended water-mask runtime placement.")
    doc.add_heading("Strengths of the Proposed Solution", level=2)
    add_bullets(
        doc,
        [
            "The pipeline exposes a coherent end-to-end workflow rather than isolated scripts.",
            "The service-oriented architecture improves maintainability and deployment flexibility.",
            "Zarr outputs make downstream access more structured and reusable.",
            "Cloud and water masks are explicit artifacts that enrich existing datasets.",
            "Stage-level timing and error reporting make operational debugging easier.",
        ],
    )
    doc.add_heading("Current Limitations", level=2)
    add_bullets(
        doc,
        [
            "Full-scene processing remains heavy because satellite products are large by nature.",
            "Precise water masking requires accelerator-backed inference for production-like scenes.",
            "Provider downloads can still be affected by external rate limits or network instability.",
            "The user interface needs continued refinement to make artifact inspection clearer.",
        ],
    )
    doc.add_heading("Possible Improvements", level=2)
    add_bullets(
        doc,
        [
            "Deploy the mask service on an OCI GPU instance while keeping the local UI and API workflow unchanged.",
            "Add a host-native MPS runner for local Mac acceleration when the dependency stack is stable.",
            "Improve UI visualization for RGB layers, cloud probability and water mask overlays.",
            "Add more stage-specific metrics, including model tile count, accelerator type and per-scene timing.",
            "Extend validation with representative water-rich, cloud-rich and multi-tile scenarios.",
        ],
    )
    doc.add_heading("Conclusion", level=2)
    add_body(
        doc,
        "The validation phase confirms the value of the modular architecture and reveals the main optimization priority. The pipeline itself can orchestrate the required stages, but precise and fast water masking needs to run where the OmniWater model can access a real accelerator.",
    )


def add_conclusion_and_refs(doc: Document) -> None:
    doc.add_heading("General Conclusion and Perspectives", level=1)
    add_body(
        doc,
        "This end-of-studies project focused on the design and implementation of a modular satellite imagery acquisition and processing pipeline. The work transformed a fragmented workflow into a service-oriented platform that can preview provider products, launch asynchronous jobs, download scenes, normalize outputs to Zarr, apply cloud and water masks and expose artifacts through a traceable lifecycle.",
    )
    add_body(
        doc,
        "The main contribution is architectural as much as technical. By separating UI, API, worker, Zarr conversion and masking responsibilities, the platform becomes easier to maintain, test, deploy and extend. This separation also makes it possible to place heavy stages on specialized infrastructure without changing the user-facing workflow.",
    )
    add_body(
        doc,
        "Future work should focus on accelerator-backed water masking, stronger UI visualization of results, broader validation across providers and tile configurations, and cloud deployment patterns that keep large raster data close to compute resources. These improvements would move the platform closer to operational use for Digital Government geospatial services.",
    )

    doc.add_heading("Webography", level=1)
    refs = [
        "[1] Oracle Corporation, official documentation and corporate information.",
        "[2] Oracle Cloud Infrastructure documentation.",
        "[3] European Space Agency, Sentinel-2 user documentation.",
        "[4] Copernicus Data Space Ecosystem documentation.",
        "[5] U.S. Geological Survey, Landsat Collection 2 and M2M API documentation.",
        "[6] Zarr Developers, Zarr format documentation and specification.",
        "[7] FastAPI documentation.",
        "[8] Streamlit documentation.",
        "[9] Rasterio, Xarray and PyProj documentation.",
        "[10] NimbusChain internal repository documentation: README, API reference, Zarr notes and runbook.",
    ]
    for ref in refs:
        add_body(doc, ref)


def main() -> None:
    doc = Document()
    setup_styles(doc)
    add_cover(doc)
    add_front_matter(doc)
    add_chapter_1(doc)
    doc.add_page_break()
    add_chapter_2(doc)
    doc.add_page_break()
    add_chapter_3(doc)
    doc.add_page_break()
    add_chapter_4(doc)
    doc.add_page_break()
    add_chapter_5(doc)
    doc.add_page_break()
    add_chapter_6(doc)
    doc.add_page_break()
    add_conclusion_and_refs(doc)
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
